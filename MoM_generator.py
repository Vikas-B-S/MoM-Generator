import streamlit as st
import google.generativeai as genai
import os
from pdfextracter import text_extractor_pdf
from docxextracter import text_extracter_docx
from imageextracter import text_extracter_image
from docx import Document
import io

# Configure the model

key=os.getenv('GOOGLE_API_KEY')
genai.configure(api_key=key)

model=genai.GenerativeModel('gemini-2.5-flash-lite')

#Upload the file in sidebar

st.sidebar.title(':orange[UPLOAD YOUR MoM NOTES HERE]')
st.sidebar.subheader('Only upload IMAGES, PDFs and DOCX files')
user_file=st.sidebar.file_uploader('Upload Your File',type=['pdf','docx','png','jpg','jpeg'])
if user_file:
  if user_file.type=='application/pdf':
    user_text=text_extractor_pdf(user_file)
  elif user_file.type=='application/vnd.openxmlformats-officedocument.wordprocessingml.document':
    user_text=text_extracter_docx(user_file)
  elif user_file.type in ['image/png','image/jpg','image/jpeg']:
    user_text=text_extracter_image(user_file)
  else: 
    st.sidebar.error('Please Upload pdf , docx or image file')


# Main page 

st.title(':blue[MINUTES OF MEETING :] :grey[AI assisted MoM generator in a standardized form from meeting notes.]')

tips = '''
### How to Use This App
- Upload your meeting notes in the sidebar (supported formats: PDF, DOCX, PNG, JPG, JPEG).  
- The AI will process your file and extract the key information.  
- Click **Generate MoM** to create a professional, structured Minutes of Meeting (MoM).  
- Download the result instantly as a Word document for sharing or record-keeping.  

**Best Practices:**
- Upload clear, legible documents for best results.  
- Ensure the meeting notes include details like attendees, agenda, and action items.  
- Review the generated MoM and make minor edits if needed before circulation.  
'''
st.write(tips)
 
if st.button('Generate MoM'):
  if user_text is None:
    st.error('Text is not generated')
  else:
    with st.spinner('Processing your data.....'):
      prompt = f"""
      You are an AI assistant specialized in writing **professional, standardized Minutes of Meeting (MoM)** 
      from uploaded content (PDF, DOCX, or image).

      ### Instructions:
      1. Carefully read and understand the uploaded meeting notes.  
      2. Extract and organize the following sections wherever possible:
        - **Meeting Title**  
        - **Date & Time**  
        - **Venue / Platform**  
        - **Attendees**  
        - **Agenda Items**  
        - **Discussion Points**  
        - **Decisions Made**  
        - **Action Items** (with responsible person and deadline, if mentioned)  

      3. Present the MoM in a clean, structured, and easy-to-read format using bullet points.  
      4. Highlight important keywords (like **decisions**, **deadlines**, and **responsible persons**) in **bold** for quick scanning.  
      5. If certain details are not available in the document, write *"Not Mentioned"*.  
      6. Ensure the output is professional and ready to be copy-pasted directly into a Word document.  

      ---

      ### Example Output Format

      **Title:** [Meeting Title]  
      **Date:** [DD-MM-YYYY]  
      **Time:** [HH:MM]  
      **Venue/Platform:** [Location/Tool]  
      **Attendees:**  
      - Person A  
      - Person B  

      **Agenda:**  
      - Item 1  
      - Item 2  

      **Discussion Points:**  
      - Key discussion point 1  
      - Key discussion point 2  

      **Decisions Made:**  
      - **Decision 1**  
      - **Decision 2**  

      **Action Items:**  
      - Task 1 → **Responsible:** Person A → **Deadline:** DD-MM-YYYY  
      - Task 2 → **Responsible:** Person B → **Deadline:** DD-MM-YYYY  

      ---

      Now, generate the MoM based strictly on the following document content:  

      {user_text}
      """

      response=model.generate_content(prompt)
      st.markdown(response.text)

      doc = Document()
      doc.add_heading("Minutes_of_Meetings", 0)
      doc.add_paragraph(response.text)

      # Save to in-memory file
      buffer = io.BytesIO()
      doc.save(buffer)
      buffer.seek(0)

      st.download_button(
      label="Download Report as Word",
      data=buffer,
      file_name="Minutes_of_Meetings.docx",
      mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
      )